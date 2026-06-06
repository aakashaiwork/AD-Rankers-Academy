@app.get("/")
async def root():
    return {
        "status": "success",
        "message": "AD Rankers Academy API Running"
    }
from fastapi import FastAPI, HTTPException, status, Depends, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, EmailStr, field_validator
from typing import List, Optional
from datetime import datetime, timedelta
from pymongo import MongoClient
from bson import ObjectId
import bcrypt
import jwt
import os
import certifi
import io
import base64
from docx import Document
from openpyxl import load_workbook
from dotenv import load_dotenv
import re

load_dotenv()

app = FastAPI()


@app.exception_handler(RequestValidationError)
async def request_validation_handler(request: Request, exc: RequestValidationError):
    if request.method == "POST" and "/api/materials" in request.url.path:
        print("[Upload Material] RequestValidationError (422):", exc.errors())
    return JSONResponse(status_code=422, content={"detail": exc.errors()})


# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# MongoDB connection
MONGO_URL = os.getenv("MONGO_URL")
DB_NAME = os.getenv("DB_NAME", "exam_app")
MONGO_TLS_INSECURE = os.getenv("MONGO_TLS_INSECURE", "false").lower() in ("1", "true", "yes")

if not MONGO_URL:
    raise RuntimeError("MONGO_URL is not set. Add it to backend/.env")

client = MongoClient(MONGO_URL)
db = client[DB_NAME]

# Collections
users_collection = db.users
categories_collection = db.categories
subjects_collection = db.subjects
materials_collection = db.materials
videos_collection = db.videos
tests_collection = db.tests
courses_collection = db.courses
attempts_collection = db.attempts
subscriptions_collection = db.subscriptions
purchases_collection = db.purchases

SECRET_KEY = os.getenv("JWT_SECRET", "your-secret-key-change-in-production")
ALGORITHM = "HS256"

# Models
class UserRegister(BaseModel):
    email: EmailStr
    password: str
    name: str
    role: str = "student"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class Category(BaseModel):
    name: str
    description: str
    icon: str

class Subject(BaseModel):
    name: str
    categoryId: str
    description: str
    icon: str

class StudyMaterial(BaseModel):
    title: str
    categoryId: str
    subjectId: str
    description: str
    fileData: str
    fileType: str
    isPaid: bool = False
    price: float = 0
    duration: int = 0  # days

    @field_validator("categoryId", "subjectId", mode="before")
    @classmethod
    def validate_object_id_fields(cls, v):
        if v is None:
            raise ValueError("cannot be empty")
        s = str(v).strip()
        if not s:
            raise ValueError("cannot be empty")
        try:
            ObjectId(s)
        except Exception:
            raise ValueError("must be a valid MongoDB ObjectId hex string")
        return s

class Video(BaseModel):
    title: str
    categoryId: str
    subjectId: str
    description: str
    youtubeUrl: str
    thumbnail: Optional[str] = None
    isPaid: bool = False
    price: float = 0
    duration: int = 0

class Question(BaseModel):
    question: str
    options: List[str]
    correctAnswer: int
    marks: int = 1
    negativeMarks: float = 0

class Test(BaseModel):
    title: str
    categoryId: str
    subjectId: str
    duration: int
    totalMarks: int
    questions: List[Question]
    isPaid: bool = False
    price: float = 0
    accessDuration: int = 0  # days

class Course(BaseModel):
    title: str
    description: str
    categoryId: str
    subjectId: str
    thumbnail: str
    accessType: str = "free"  # free | subscription | paid
    isPaid: bool = False
    price: float = 0
    duration: int = 0  # days
    videoIds: List[str] = []
    materialIds: List[str] = []
    testIds: List[str] = []

class CourseUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    categoryId: Optional[str] = None
    subjectId: Optional[str] = None
    thumbnail: Optional[str] = None
    isPaid: Optional[bool] = None
    price: Optional[float] = None
    duration: Optional[int] = None
    accessType: Optional[str] = None
    videoIds: Optional[List[str]] = None
    materialIds: Optional[List[str]] = None
    testIds: Optional[List[str]] = None

class TestSubmission(BaseModel):
    testId: str
    answers: List[int]
    timeTaken: int

class BulkTestUpload(BaseModel):
    fileData: str  # base64
    fileType: str  # docx or xlsx
    title: str
    categoryId: str
    subjectId: str
    duration: int
    isPaid: bool = False
    price: float = 0
    accessDuration: int = 0


class SubscribeRequest(BaseModel):
    days: int = 30
    type: str = "subscription"


class AdminGrantSubscriptionRequest(BaseModel):
    userId: str
    days: int = 30
    type: str = "subscription"


class PurchaseRequest(BaseModel):
    days: int = 30


class AdminGrantPurchaseRequest(BaseModel):
    userId: str
    days: int = 30

# Helper functions
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "role": role,
        "exp": datetime.utcnow() + timedelta(days=7)
    }
    return jwt.encode(payload, SECRET_KEY, algorithm=ALGORITHM)

def verify_token(token: str):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return payload
    except:
        raise HTTPException(status_code=401, detail="Invalid token")

def check_user_access(user_id: str, item_type: str, item_id: str, is_paid: bool) -> bool:
    """Check if user has access to paid content"""
    if not is_paid:
        return True
    
    # Check active subscription
    subscription = subscriptions_collection.find_one({
        "userId": user_id,
        "isActive": True,
        "endDate": {"$gt": datetime.utcnow()}
    })
    if subscription:
        return True
    
    # Check individual purchase
    purchase = purchases_collection.find_one({
        "userId": user_id,
        "itemType": item_type,
        "itemId": item_id,
        "isActive": True,
        "expiryDate": {"$gt": datetime.utcnow()}
    })
    return purchase is not None


def has_active_subscription(user_id: str) -> bool:
    subscription = subscriptions_collection.find_one({
        "userId": user_id,
        "isActive": True,
        "endDate": {"$gt": datetime.utcnow()}
    })
    return subscription is not None


def has_active_purchase(user_id: str, item_type: str, item_id: str) -> bool:
    purchase = purchases_collection.find_one({
        "userId": user_id,
        "itemType": item_type,
        "itemId": item_id,
        "isActive": True,
        "expiryDate": {"$gt": datetime.utcnow()}
    })
    return purchase is not None


def course_access_type(course: dict) -> str:
    at = course.get("accessType")
    if at in ["free", "subscription", "paid"]:
        return at
    return "paid" if course.get("isPaid", False) else "free"


def check_course_access(user_id: str, course: dict) -> bool:
    at = course_access_type(course)
    if at == "free":
        return True
    if at == "subscription":
        return has_active_subscription(user_id)
    return has_active_purchase(user_id, "course", str(course.get("_id")))

def parse_excel_test(file_data: str) -> List[Question]:
    """Parse Excel file to extract questions"""
    try:
        # Decode base64
        file_bytes = base64.b64decode(file_data.split(',')[1] if ',' in file_data else file_data)
        wb = load_workbook(io.BytesIO(file_bytes))
        ws = wb.active
        
        questions = []
        # Skip header row, start from row 2
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[0]:  # Skip empty rows
                continue
            
            question = Question(
                question=str(row[0]),
                options=[str(row[1]), str(row[2]), str(row[3]), str(row[4])],
                correctAnswer=int(row[5]) - 1,  # Convert 1-4 to 0-3
                marks=int(row[6]) if len(row) > 6 and row[6] else 1,
                negativeMarks=float(row[7]) if len(row) > 7 and row[7] is not None else 0
            )
            questions.append(question)
        
        return questions
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error parsing Excel: {str(e)}")

def parse_word_test(file_data: str) -> List[Question]:
    """Parse Word file to extract questions"""
    try:
        file_bytes = base64.b64decode(file_data.split(',')[1] if ',' in file_data else file_data)
        doc = Document(io.BytesIO(file_bytes))

        def norm(s: str) -> str:
            return ''.join(ch for ch in (s or '').strip().lower().replace(' ', '_') if ch.isalnum() or ch == '_')

        def cell_text(cell) -> str:
            return (cell.text or '').strip()

        def parse_float_safe(s: str) -> float:
            try:
                return float(str(s).strip())
            except Exception:
                return 0

        def parse_int_safe(s: str, default: int = 0) -> int:
            try:
                return int(float(str(s).strip()))
            except Exception:
                return default

        def parse_questions_from_tables() -> List[Question]:
            out: List[Question] = []
            current = None

            def flush():
                nonlocal current
                if not current:
                    return
                if not current.get('question'):
                    current = None
                    return
                if not current.get('options'):
                    current = None
                    return
                if current.get('correctAnswer') is None:
                    current['correctAnswer'] = 0
                out.append(
                    Question(
                        question=current['question'],
                        options=current['options'],
                        correctAnswer=int(current['correctAnswer']),
                        marks=int(current.get('marks', 1) or 1),
                        negativeMarks=float(current.get('negativeMarks', 0) or 0),
                    )
                )
                current = None

            for table in doc.tables:
                for row in table.rows:
                    if not row.cells:
                        continue
                    key = norm(cell_text(row.cells[0]))
                    if not key:
                        continue

                    if key == 'question':
                        flush()
                        q_text = cell_text(row.cells[1]) if len(row.cells) > 1 else ''
                        current = {
                            'question': q_text,
                            'options': [],
                            'correctAnswer': None,
                            'marks': 1,
                            'negativeMarks': 0,
                        }
                        continue

                    if current is None:
                        continue

                    if key == 'option':
                        opt_text = cell_text(row.cells[1]) if len(row.cells) > 1 else ''
                        correctness = cell_text(row.cells[2]) if len(row.cells) > 2 else ''
                        if opt_text:
                            idx = len(current['options'])
                            current['options'].append(opt_text)
                            if norm(correctness) in ['correct', 'true', 'yes']:
                                current['correctAnswer'] = idx
                        continue

                    if key == 'answer':
                        ans_text = cell_text(row.cells[1]) if len(row.cells) > 1 else ''
                        if ans_text:
                            # best-effort: map to option if exists, otherwise add as a single option
                            found = False
                            for i, opt in enumerate(current['options']):
                                if opt.strip().lower() == ans_text.strip().lower():
                                    current['correctAnswer'] = i
                                    found = True
                                    break
                            if not found:
                                current['options'] = [ans_text]
                                current['correctAnswer'] = 0
                        continue

                    if key == 'marks':
                        pos = cell_text(row.cells[1]) if len(row.cells) > 1 else ''
                        neg = cell_text(row.cells[2]) if len(row.cells) > 2 else ''
                        if pos:
                            current['marks'] = parse_int_safe(pos, 1)
                        if neg:
                            current['negativeMarks'] = parse_float_safe(neg)
                        continue

                    # ignore: type, solution, etc.

            flush()
            return out

        # Prefer table format when available
        if getattr(doc, 'tables', None) and len(doc.tables) > 0:
            table_questions = parse_questions_from_tables()
            if table_questions:
                return table_questions
        
        questions = []
        current_question = None
        options = []
        correct_answer = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            q_match = re.match(r'^(?:Q\s*\.?\s*\d*\s*[\).\-:]?\s*|\d+\s*[\).\-:]\s*)(.+)$', text, flags=re.IGNORECASE)
            opt_match = re.match(r'^([A-Da-d])\s*[\).\-:]\s*(.+)$', text)

            # Question starts with Q. / Q1. / 1. / 1) / 1- etc.
            if q_match:
                if current_question and options:
                    questions.append(Question(
                        question=current_question,
                        options=options,
                        correctAnswer=correct_answer,
                        marks=1
                    ))
                current_question = q_match.group(1).strip()
                options = []
                correct_answer = 0
            
            # Options start with A) / A. / A- / etc.
            elif opt_match:
                option_text = opt_match.group(2).strip()
                if '*' in option_text or '(correct)' in option_text.lower():
                    correct_answer = len(options)
                    option_text = option_text.replace('*', '').replace('(correct)', '').replace('(Correct)', '').strip()
                options.append(option_text)
        
        # Add last question
        if current_question and options:
            questions.append(Question(
                question=current_question,
                options=options,
                correctAnswer=correct_answer,
                marks=1,
                negativeMarks=0
            ))
        
        return questions
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error parsing Word: {str(e)}")


@app.post("/api/subscribe")
async def subscribe(req: SubscribeRequest, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    days = req.days if req.days and req.days > 0 else 30
    end = datetime.utcnow() + timedelta(days=days)

    subscriptions_collection.update_many({"userId": payload["user_id"], "isActive": True}, {"$set": {"isActive": False}})
    doc = {
        "userId": payload["user_id"],
        "type": req.type,
        "startDate": datetime.utcnow(),
        "endDate": end,
        "isActive": True,
        "createdAt": datetime.utcnow(),
    }
    subscriptions_collection.insert_one(doc)
    return {"message": "Subscription activated", "endDate": end.isoformat()}


@app.post("/api/purchase/course/{course_id}")
async def purchase_course(course_id: str, req: PurchaseRequest, token: str = None):
    # For Option A, we allow a user to create their own purchase record.
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    course = courses_collection.find_one({"_id": ObjectId(course_id)})
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")

    if course_access_type(course) != "paid":
        raise HTTPException(status_code=400, detail="This course does not require individual purchase")

    days = req.days if req.days and req.days > 0 else int(course.get("duration", 0) or 30)
    expiry = datetime.utcnow() + timedelta(days=days)

    purchases_collection.update_many(
        {"userId": payload["user_id"], "itemType": "course", "itemId": course_id, "isActive": True},
        {"$set": {"isActive": False}},
    )
    doc = {
        "userId": payload["user_id"],
        "itemType": "course",
        "itemId": course_id,
        "isActive": True,
        "purchaseDate": datetime.utcnow(),
        "expiryDate": expiry,
        "createdAt": datetime.utcnow(),
    }
    purchases_collection.insert_one(doc)
    return {"message": "Course purchased", "expiryDate": expiry.isoformat()}


@app.post("/api/admin/grant-subscription")
async def admin_grant_subscription(req: AdminGrantSubscriptionRequest, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        ObjectId(req.userId)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid user id")

    days = req.days if req.days and req.days > 0 else 30
    end = datetime.utcnow() + timedelta(days=days)

    subscriptions_collection.update_many({"userId": req.userId, "isActive": True}, {"$set": {"isActive": False}})
    doc = {
        "userId": req.userId,
        "type": req.type,
        "startDate": datetime.utcnow(),
        "endDate": end,
        "isActive": True,
        "createdAt": datetime.utcnow(),
        "grantedBy": payload["user_id"],
    }
    subscriptions_collection.insert_one(doc)
    return {"message": "Subscription granted", "endDate": end.isoformat()}


@app.post("/api/admin/grant-course-purchase/{course_id}")
async def admin_grant_course_purchase(course_id: str, req: AdminGrantPurchaseRequest, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    course = courses_collection.find_one({"_id": ObjectId(course_id)})
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")

    if course_access_type(course) != "paid":
        raise HTTPException(status_code=400, detail="This course does not require individual purchase")

    try:
        ObjectId(req.userId)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid user id")

    days = req.days if req.days and req.days > 0 else int(course.get("duration", 0) or 30)
    expiry = datetime.utcnow() + timedelta(days=days)

    purchases_collection.update_many(
        {"userId": req.userId, "itemType": "course", "itemId": course_id, "isActive": True},
        {"$set": {"isActive": False}},
    )
    doc = {
        "userId": req.userId,
        "itemType": "course",
        "itemId": course_id,
        "isActive": True,
        "purchaseDate": datetime.utcnow(),
        "expiryDate": expiry,
        "createdAt": datetime.utcnow(),
        "grantedBy": payload["user_id"],
    }
    purchases_collection.insert_one(doc)
    return {"message": "Course purchase granted", "expiryDate": expiry.isoformat()}

# Seed default data
@app.on_event("startup")
async def startup_event():
    # Create default categories
    if categories_collection.count_documents({}) == 0:
        default_categories = [
            {"name": "SSC", "description": "Staff Selection Commission", "icon": "briefcase"},
            {"name": "BANK", "description": "Banking Exams", "icon": "business"},
            {"name": "GPSC", "description": "Gujarat Public Service Commission", "icon": "school"},
            {"name": "UPSC", "description": "Union Public Service Commission", "icon": "shield"},
            {"name": "Other", "description": "Other Competitive Exams", "icon": "list"}
        ]
        result = categories_collection.insert_many(default_categories)
        category_ids = {cat["name"]: str(id) for cat, id in zip(default_categories, result.inserted_ids)}
    else:
        categories = list(categories_collection.find())
        category_ids = {cat["name"]: str(cat["_id"]) for cat in categories}
    
    # Create default subjects
    if subjects_collection.count_documents({}) == 0:
        default_subjects = [
            {"name": "Mathematics", "categoryId": category_ids.get("SSC", ""), "icon": "calculator", "description": "Math and quantitative aptitude"},
            {"name": "Science", "categoryId": category_ids.get("SSC", ""), "icon": "flask", "description": "Physics, Chemistry, Biology"},
            {"name": "English", "categoryId": category_ids.get("SSC", ""), "icon": "book", "description": "Grammar, vocabulary, comprehension"},
            {"name": "General Knowledge", "categoryId": category_ids.get("UPSC", ""), "icon": "globe", "description": "Current affairs and GK"},
            {"name": "Reasoning", "categoryId": category_ids.get("SSC", ""), "icon": "brain", "description": "Logical and analytical reasoning"}
        ]
        subjects_collection.insert_many(default_subjects)
    
    # Create default admin
    if users_collection.count_documents({"role": "admin"}) == 0:
        admin_user = {
            "email": "admin@exam.com",
            "password": hash_password("admin123"),
            "name": "Admin User",
            "role": "admin",
            "createdAt": datetime.utcnow(),
            "progress": {"completedTests": [], "watchedVideos": [], "totalScore": 0}
        }
        users_collection.insert_one(admin_user)

# Routes
@app.get("/api/health")
async def health():
    return {"status": "ok", "message": "Exam App API is running"}

# Auth Routes
@app.post("/api/register")
async def register(user: UserRegister):
    if users_collection.find_one({"email": user.email}):
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_doc = {
        "email": user.email,
        "password": hash_password(user.password),
        "name": user.name,
        "role": user.role if user.role in ["student", "admin"] else "student",
        "createdAt": datetime.utcnow(),
        "progress": {"completedTests": [], "watchedVideos": [], "totalScore": 0}
    }
    
    result = users_collection.insert_one(user_doc)
    user_id = str(result.inserted_id)
    token = create_token(user_id, user.email, user_doc["role"])
    
    return {
        "token": token,
        "user": {
            "id": user_id,
            "email": user.email,
            "name": user.name,
            "role": user_doc["role"]
        }
    }

@app.post("/api/login")
async def login(credentials: UserLogin):
    user = users_collection.find_one({"email": credentials.email})
    
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    user_id = str(user["_id"])
    token = create_token(user_id, user["email"], user["role"])
    
    return {
        "token": token,
        "user": {
            "id": user_id,
            "email": user["email"],
            "name": user["name"],
            "role": user["role"]
        }
    }

# Category Routes
@app.get("/api/categories")
async def get_categories():
    categories = list(categories_collection.find())
    for cat in categories:
        cat["_id"] = str(cat["_id"])
    return {"categories": categories}

@app.post("/api/categories")
async def create_category(category: Category, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    cat_doc = {**category.dict(), "createdAt": datetime.utcnow()}
    result = categories_collection.insert_one(cat_doc)
    return {"id": str(result.inserted_id), "message": "Category created successfully"}


@app.delete("/api/categories/{category_id}")
async def delete_category(category_id: str, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        oid = ObjectId(category_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid category id")

    result = categories_collection.delete_one({"_id": oid})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Category not found")

    return {"message": "Category deleted successfully"}

# Subject Routes
@app.get("/api/subjects")
async def get_subjects(categoryId: Optional[str] = None):
    query = {}
    if categoryId:
        query["categoryId"] = categoryId
    
    subjects = list(subjects_collection.find(query))
    for subject in subjects:
        subject["_id"] = str(subject["_id"])
        if subject.get("categoryId") is not None:
            subject["categoryId"] = str(subject["categoryId"])
    return {"subjects": subjects}

@app.post("/api/subjects")
async def create_subject(subject: Subject, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    name_stripped = subject.name.strip()
    if not name_stripped:
        raise HTTPException(status_code=400, detail="Subject name cannot be empty")

    name_lower = name_stripped.lower()
    for existing in subjects_collection.find({"categoryId": subject.categoryId}):
        if (existing.get("name") or "").strip().lower() == name_lower:
            raise HTTPException(status_code=400, detail="Subject already exists in this category")

    subj_payload = subject.dict()
    subj_payload["name"] = name_stripped
    subj_doc = {**subj_payload, "createdAt": datetime.utcnow()}
    result = subjects_collection.insert_one(subj_doc)
    return {"id": str(result.inserted_id), "message": "Subject created successfully"}

@app.delete("/api/subjects/{subject_id}")
async def delete_subject(subject_id: str, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        oid = ObjectId(subject_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid subject id")

    result = subjects_collection.delete_one({"_id": oid})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Subject not found")

    return {"message": "Subject deleted successfully"}

# Study Material Routes
@app.get("/api/materials")
async def get_materials(categoryId: Optional[str] = None, subjectId: Optional[str] = None, token: Optional[str] = None):
    query = {}
    if categoryId:
        query["categoryId"] = categoryId
    if subjectId:
        query["subjectId"] = subjectId
    
    materials = list(materials_collection.find(query).sort("createdAt", -1))
    
    # Filter based on user access for paid content
    if token:
        try:
            payload = verify_token(token)
            user_id = payload["user_id"]
            filtered_materials = []
            for material in materials:
                material["_id"] = str(material["_id"])
                material["hasAccess"] = check_user_access(user_id, "material", str(material["_id"]), material.get("isPaid", False))
                if not material.get("isPaid", False) or material["hasAccess"]:
                    filtered_materials.append(material)
                else:
                    # Show preview for paid content
                    material["fileData"] = ""
                    filtered_materials.append(material)
            return {"materials": filtered_materials}
        except:
            pass
    
    for material in materials:
        material["_id"] = str(material["_id"])
        material["hasAccess"] = not material.get("isPaid", False)
        if material.get("isPaid", False):
            material["fileData"] = ""
    
    return {"materials": materials}

@app.get("/api/materials/{material_id}")
async def get_material(material_id: str, token: Optional[str] = None):
    material = materials_collection.find_one({"_id": ObjectId(material_id)})
    if not material:
        raise HTTPException(status_code=404, detail="Material not found")
    
    material["_id"] = str(material["_id"])
    
    # Check access for paid content
    if material.get("isPaid", False):
        if not token:
            raise HTTPException(status_code=401, detail="Authentication required for paid content")
        
        payload = verify_token(token)
        has_access = check_user_access(payload["user_id"], "material", material_id, True)
        
        if not has_access:
            raise HTTPException(status_code=403, detail="Purchase required to access this content")
    
    return material

@app.post("/api/materials")
async def create_material(material: StudyMaterial, token: str = None):
    print(
        "[Upload Material] route entered | format: JSON (application/json); fileData is a string (raw base64 or data URL with base64)"
    )
    print(
        "[Upload Material] fields:",
        "title=",
        repr(material.title[:80] if material.title else ""),
        "categoryId=",
        material.categoryId,
        "subjectId=",
        material.subjectId,
        "fileType=",
        material.fileType,
        "fileData_len=",
        len(material.fileData or ""),
        "fileData_prefix=",
        repr((material.fileData or "")[:80]),
    )

    if not token:
        print("[Upload Material] error: missing token query param")
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        print("[Upload Material] error: forbidden, role=", payload.get("role"))
        raise HTTPException(status_code=403, detail="Admin access required")

    material_doc = {
        **material.model_dump(),
        "uploadedBy": payload["user_id"],
        "createdAt": datetime.utcnow(),
    }

    try:
        result = materials_collection.insert_one(material_doc)
    except Exception as e:
        print("[Upload Material] MongoDB insert_one failed:", repr(e))
        raise HTTPException(
            status_code=500,
            detail=f"Database error while saving material: {str(e)}",
        )

    print("[Upload Material] success, inserted_id=", str(result.inserted_id))
    return {"id": str(result.inserted_id), "message": "Material uploaded successfully"}

@app.delete("/api/materials/{material_id}")
async def delete_material(material_id: str, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    result = materials_collection.delete_one({"_id": ObjectId(material_id)})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Material not found")
    return {"message": "Material deleted successfully"}

# Video Routes
@app.get("/api/videos")
async def get_videos(categoryId: Optional[str] = None, subjectId: Optional[str] = None, token: Optional[str] = None):
    query = {}
    if categoryId:
        query["categoryId"] = categoryId
    if subjectId:
        query["subjectId"] = subjectId
    
    videos = list(videos_collection.find(query).sort("createdAt", -1))
    
    if token:
        try:
            payload = verify_token(token)
            user_id = payload["user_id"]
            for video in videos:
                video["_id"] = str(video["_id"])
                video["hasAccess"] = check_user_access(user_id, "video", str(video["_id"]), video.get("isPaid", False))
        except:
            pass
    else:
        for video in videos:
            video["_id"] = str(video["_id"])
            video["hasAccess"] = not video.get("isPaid", False)
    
    return {"videos": videos}

@app.post("/api/videos")
async def create_video(video: Video, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    video_doc = {
        **video.dict(),
        "uploadedBy": payload["user_id"],
        "createdAt": datetime.utcnow()
    }
    
    result = videos_collection.insert_one(video_doc)
    return {"id": str(result.inserted_id), "message": "Video added successfully"}

@app.delete("/api/videos/{video_id}")
async def delete_video(video_id: str, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    result = videos_collection.delete_one({"_id": ObjectId(video_id)})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Video not found")
    return {"message": "Video deleted successfully"}

# Test Routes
@app.get("/api/tests")
async def get_tests(categoryId: Optional[str] = None, subjectId: Optional[str] = None, token: Optional[str] = None):
    query = {}
    if categoryId:
        query["categoryId"] = categoryId
    if subjectId:
        query["subjectId"] = subjectId
    
    tests = list(tests_collection.find(query).sort("createdAt", -1))
    
    for test in tests:
        test["_id"] = str(test["_id"])
        # Don't send correct answers
        for q in test["questions"]:
            q.pop("correctAnswer", None)
        
        if token:
            try:
                payload = verify_token(token)
                test["hasAccess"] = check_user_access(payload["user_id"], "test", str(test["_id"]), test.get("isPaid", False))
            except:
                test["hasAccess"] = not test.get("isPaid", False)
        else:
            test["hasAccess"] = not test.get("isPaid", False)
    
    return {"tests": tests}

@app.get("/api/tests/{test_id}")
async def get_test(test_id: str, token: Optional[str] = None):
    test = tests_collection.find_one({"_id": ObjectId(test_id)})
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")
    
    test["_id"] = str(test["_id"])
    
    # Check access for paid content
    if test.get("isPaid", False):
        if not token:
            raise HTTPException(status_code=401, detail="Authentication required for paid content")
        
        payload = verify_token(token)
        has_access = check_user_access(payload["user_id"], "test", test_id, True)
        
        if not has_access:
            raise HTTPException(status_code=403, detail="Purchase required to access this test")
    
    # Don't send correct answers
    for q in test["questions"]:
        q.pop("correctAnswer", None)
    return test

@app.post("/api/tests")
async def create_test(test: Test, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    test_doc = {
        **test.dict(),
        "createdBy": payload["user_id"],
        "createdAt": datetime.utcnow()
    }
    
    result = tests_collection.insert_one(test_doc)
    return {"id": str(result.inserted_id), "message": "Test created successfully"}

@app.post("/api/tests/bulk-upload")
async def bulk_upload_test(upload: BulkTestUpload, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Parse file based on type
    if upload.fileType == "xlsx":
        questions = parse_excel_test(upload.fileData)
    elif upload.fileType == "docx":
        questions = parse_word_test(upload.fileData)
    else:
        raise HTTPException(status_code=400, detail="Invalid file type. Use xlsx or docx")
    
    if not questions:
        raise HTTPException(status_code=400, detail="No questions found in file")
    
    # Calculate total marks
    total_marks = sum(q.marks for q in questions)
    
    test_doc = {
        "title": upload.title,
        "categoryId": upload.categoryId,
        "subjectId": upload.subjectId,
        "duration": upload.duration,
        "totalMarks": total_marks,
        "questions": [q.dict() for q in questions],
        "isPaid": upload.isPaid,
        "price": upload.price,
        "accessDuration": upload.accessDuration,
        "createdBy": payload["user_id"],
        "createdAt": datetime.utcnow()
    }
    
    result = tests_collection.insert_one(test_doc)
    return {
        "id": str(result.inserted_id),
        "message": f"Test created successfully with {len(questions)} questions",
        "questionCount": len(questions),
        "totalMarks": total_marks
    }


@app.get("/api/admin/templates/test-docx")
async def download_test_docx_template(token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    doc = Document()
    doc.add_heading("MCQ Bulk Upload - DOCX Template", level=1)
    doc.add_paragraph("Use either the table format (recommended) or the paragraph format.")

    doc.add_heading("Table Format (Recommended)", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    hdr[2].text = "Correct (for Option row) / Negative (for Marks row)"

    def add_row(a: str, b: str = "", c: str = ""):
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    add_row("Question", "What is 2 + 2?")
    add_row("Option", "3", "incorrect")
    add_row("Option", "4", "correct")
    add_row("Option", "5", "incorrect")
    add_row("Option", "22", "incorrect")
    add_row("Marks", "1", "0.25")

    doc.add_paragraph(" ")
    doc.add_heading("Paragraph Format (Alternative)", level=2)
    doc.add_paragraph("Q1. What is 2 + 2?")
    doc.add_paragraph("A) 3")
    doc.add_paragraph("B) 4*")
    doc.add_paragraph("C) 5")
    doc.add_paragraph("D) 22")
    doc.add_paragraph("(Use * to mark the correct option)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="mcq_upload_template.docx"'
        },
    )

@app.post("/api/tests/submit")
async def submit_test(submission: TestSubmission, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    
    test = tests_collection.find_one({"_id": ObjectId(submission.testId)})
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")
    
    # Check access for paid tests
    if test.get("isPaid", False):
        has_access = check_user_access(payload["user_id"], "test", submission.testId, True)
        if not has_access:
            raise HTTPException(status_code=403, detail="Purchase required to take this test")
    
    # Calculate score
    score = 0
    for i, answer in enumerate(submission.answers):
        if i < len(test["questions"]):
            q = test["questions"][i]
            if answer == q.get("correctAnswer"):
                score += q.get("marks", 1)
            elif answer is not None and int(answer) != -1:
                score -= float(q.get("negativeMarks", 0) or 0)
    
    attempt_doc = {
        "testId": submission.testId,
        "userId": payload["user_id"],
        "answers": submission.answers,
        "score": score,
        "totalMarks": test["totalMarks"],
        "timeTaken": submission.timeTaken,
        "attemptedAt": datetime.utcnow()
    }
    
    result = attempts_collection.insert_one(attempt_doc)
    
    users_collection.update_one(
        {"_id": ObjectId(payload["user_id"])},
        {
            "$push": {"progress.completedTests": submission.testId},
            "$inc": {"progress.totalScore": score}
        }
    )
    
    return {
        "attemptId": str(result.inserted_id),
        "score": score,
        "totalMarks": test["totalMarks"],
        "percentage": (score / test["totalMarks"] * 100) if test["totalMarks"] > 0 else 0
    }

# Course Routes
@app.get("/api/courses")
async def get_courses(categoryId: Optional[str] = None, subjectId: Optional[str] = None, token: Optional[str] = None):
    query = {}
    if categoryId:
        query["categoryId"] = categoryId
    if subjectId:
        query["subjectId"] = subjectId
    
    courses = list(courses_collection.find(query).sort("createdAt", -1))
    
    for course in courses:
        course["_id"] = str(course["_id"])
        course["accessType"] = course_access_type(course)
        if token:
            try:
                payload = verify_token(token)
                course["hasAccess"] = check_course_access(payload["user_id"], course)
            except:
                course["hasAccess"] = course_access_type(course) == "free"
        else:
            course["hasAccess"] = course_access_type(course) == "free"
    
    return {"courses": courses}

@app.get("/api/courses/{course_id}")
async def get_course(course_id: str, token: Optional[str] = None):
    course = courses_collection.find_one({"_id": ObjectId(course_id)})
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")
    
    course["_id"] = str(course["_id"])
    
    course["accessType"] = course_access_type(course)

    # Check access (3-tier)
    has_access = course_access_type(course) == "free"
    if token:
        payload = verify_token(token)
        if payload.get("role") == "admin":
            has_access = True
        else:
            has_access = check_course_access(payload["user_id"], course)
    course["hasAccess"] = has_access

    if not has_access:
        return course

    # Populate course items
    videos = []
    for vid_id in course.get("videoIds", []):
        video = videos_collection.find_one({"_id": ObjectId(vid_id)})
        if video:
            video["_id"] = str(video["_id"])
            videos.append(video)
    
    materials = []
    for mat_id in course.get("materialIds", []):
        material = materials_collection.find_one({"_id": ObjectId(mat_id)})
        if material:
            material["_id"] = str(material["_id"])
            materials.append(material)
    
    tests = []
    for test_id in course.get("testIds", []):
        test = tests_collection.find_one({"_id": ObjectId(test_id)})
        if test:
            test["_id"] = str(test["_id"])
            tests.append(test)
    
    course["videos"] = videos
    course["materials"] = materials
    course["tests"] = tests
    
    return course


@app.delete("/api/courses/{course_id}")
async def delete_course(course_id: str, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        oid = ObjectId(course_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid course id")

    result = courses_collection.delete_one({"_id": oid})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Course not found")

    return {"message": "Course deleted successfully"}

@app.post("/api/courses")
async def create_course(course: Course, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    course_doc = {
        **course.dict(),
        "createdBy": payload["user_id"],
        "createdAt": datetime.utcnow()
    }
    
    result = courses_collection.insert_one(course_doc)
    return {"id": str(result.inserted_id), "message": "Course created successfully"}


@app.put("/api/courses/{course_id}")
async def update_course(course_id: str, update: CourseUpdate, token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")

    try:
        oid = ObjectId(course_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid course id")

    course = courses_collection.find_one({"_id": oid})
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")

    data = update.dict(exclude_unset=True)

    # Normalize accessType/isPaid
    if "accessType" in data and data["accessType"] is not None:
        at = str(data["accessType"]).lower().strip()
        if at not in ["free", "subscription", "paid"]:
            raise HTTPException(status_code=400, detail="Invalid accessType. Use free/subscription/paid")
        data["accessType"] = at
        data["isPaid"] = at == "paid"

    if "isPaid" in data and "accessType" not in data and data["isPaid"] is not None:
        data["accessType"] = "paid" if bool(data["isPaid"]) else "free"

    data["updatedAt"] = datetime.utcnow()
    data["updatedBy"] = payload["user_id"]

    result = courses_collection.update_one({"_id": oid}, {"$set": data})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Course not found")

    return {"message": "Course updated successfully"}

# User Profile
@app.get("/api/profile")
async def get_profile(token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    user = users_collection.find_one({"_id": ObjectId(payload["user_id"])})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    attempts = list(attempts_collection.find({"userId": payload["user_id"]}).sort("attemptedAt", -1).limit(10))
    for attempt in attempts:
        attempt["_id"] = str(attempt["_id"])
    
    # Get subscription info
    subscription = subscriptions_collection.find_one({
        "userId": payload["user_id"],
        "isActive": True,
        "endDate": {"$gt": datetime.utcnow()}
    })
    
    subscription_info = None
    if subscription:
        subscription_info = {
            "type": subscription["type"],
            "endDate": subscription["endDate"].isoformat(),
            "isActive": True
        }
    
    return {
        "user": {
            "id": str(user["_id"]),
            "email": user["email"],
            "name": user["name"],
            "role": user["role"],
            "progress": user.get("progress", {})
        },
        "recentAttempts": attempts,
        "subscription": subscription_info
    }

# Leaderboard
@app.get("/api/leaderboard")
async def get_leaderboard():
    users = list(users_collection.find({"role": "student"}).sort("progress.totalScore", -1).limit(20))
    leaderboard = []
    for i, user in enumerate(users):
        leaderboard.append({
            "rank": i + 1,
            "name": user["name"],
            "totalScore": user.get("progress", {}).get("totalScore", 0),
            "testsCompleted": len(user.get("progress", {}).get("completedTests", []))
        })
    return {"leaderboard": leaderboard}

# Admin Stats
@app.get("/api/admin/stats")
async def get_admin_stats(token: str = None):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")
    
    payload = verify_token(token)
    if payload["role"] != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    stats = {
        "totalStudents": users_collection.count_documents({"role": "student"}),
        "totalTests": tests_collection.count_documents({}),
        "totalMaterials": materials_collection.count_documents({}),
        "totalVideos": videos_collection.count_documents({}),
        "totalCourses": courses_collection.count_documents({}),
        "totalAttempts": attempts_collection.count_documents({}),
        "totalCategories": categories_collection.count_documents({})
    }
    
    return stats

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
